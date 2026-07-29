from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = ROOT / 'PLANO_ESTUDOS_CONCURSO_245_ITENS.md'
DOCX_PATH = ROOT / 'PLANO_ESTUDOS_CONCURSO_245_ITENS.docx'


def set_run_font(run, name: str = 'Calibri', size: float = 11, color: str | None = None, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{margin}'))
        if node is None:
            node = OxmlElement(f'w:{margin}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width_dxa: int) -> None:
    cell.width = Inches(width_dxa / 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in('w:tcW')
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths_dxa)))
    tbl_w.set(qn('w:type'), 'dxa')

    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in('w:shd')
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    field_begin = OxmlElement('w:fldChar')
    field_begin.set(qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText')
    instruction.set(qn('xml:space'), 'preserve')
    instruction.text = ' PAGE '
    field_end = OxmlElement('w:fldChar')
    field_end.set(qn('w:fldCharType'), 'end')
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)
    set_run_font(run, size=9, color='6B7280')


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ('Heading 1', 16, '2E74B5', 18, 10),
        ('Heading 2', 13, '2E74B5', 14, 7),
        ('Heading 3', 12, '1F4D78', 10, 5),
    ]:
        style = styles[name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if 'Checklist' not in styles:
        checklist = styles.add_style('Checklist', WD_STYLE_TYPE.PARAGRAPH)
    else:
        checklist = styles['Checklist']
    checklist.base_style = styles['Normal']
    checklist.font.name = 'Calibri'
    checklist._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    checklist._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    checklist.font.size = Pt(10.5)
    checklist.paragraph_format.left_indent = Inches(0.16)
    checklist.paragraph_format.first_line_indent = Inches(-0.16)
    checklist.paragraph_format.space_after = Pt(3)
    checklist.paragraph_format.line_spacing = 1.15

    if 'Compact Metadata' not in styles:
        metadata = styles.add_style('Compact Metadata', WD_STYLE_TYPE.PARAGRAPH)
    else:
        metadata = styles['Compact Metadata']
    metadata.base_style = styles['Normal']
    metadata.font.name = 'Calibri'
    metadata.font.size = Pt(10)
    metadata.font.color.rgb = RGBColor.from_string('4B5563')
    metadata.paragraph_format.space_after = Pt(4)


def add_inline_text(paragraph, text: str) -> None:
    pattern = re.compile(r'\*\*([^*]+)\*\*')
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run)
        run = paragraph.add_run(match.group(1))
        set_run_font(run, color='0B2545', bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run)


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=24, color='0B2545', bold=True)


def add_summary_table(document: Document, lines: list[str], index: int) -> int:
    rows = []
    while index < len(lines) and lines[index].startswith('|'):
        cells = [cell.strip() for cell in lines[index].strip('|').split('|')]
        if cells and not all(set(cell) <= {'-', ':'} for cell in cells):
            rows.append(cells)
        index += 1

    if not rows:
        return index

    table = document.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [7200, 2160])
    table.style = 'Table Grid'
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data[:2]):
            cell = table.cell(row_index, col_index)
            cell.text = ''
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=10, color='0B2545' if row_index == 0 else '111827', bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, 'E8EEF5')
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return index


def build_document() -> None:
    lines = MARKDOWN_PATH.read_text(encoding='utf-8').splitlines()
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run('PLANO DE ESTUDOS | CONCURSO PÚBLICO')
    set_run_font(header_run, size=9, color='6B7280', bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run('Página ')
    set_run_font(footer_run, size=9, color='6B7280')
    add_page_number(footer)

    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith('# '):
            add_title(document, line[2:].strip())
        elif line.startswith('## '):
            paragraph = document.add_paragraph(line[3:].strip(), style='Heading 1')
            paragraph.paragraph_format.keep_with_next = True
        elif line.startswith('### '):
            document.add_paragraph(line[4:].strip(), style='Heading 2')
        elif line.startswith('|'):
            index = add_summary_table(document, lines, index)
            continue
        elif re.match(r'- \[ \] \*\*[^*]+\*\* ', line):
            match = re.match(r'- \[ \] \*\*([^*]+)\*\* (.*)', line)
            paragraph = document.add_paragraph(style='Checklist')
            checkbox = paragraph.add_run('[ ] ')
            set_run_font(checkbox, size=10.5, color='374151')
            code = paragraph.add_run(match.group(1))
            set_run_font(code, size=10.5, color='0B2545', bold=True)
            separator = paragraph.add_run('  ')
            set_run_font(separator, size=10.5)
            add_inline_text(paragraph, match.group(2))
        elif re.match(r'\d+\. ', line):
            paragraph = document.add_paragraph(style='List Number')
            add_inline_text(paragraph, re.sub(r'^\d+\. ', '', line))
        elif line.startswith('- '):
            paragraph = document.add_paragraph(style='List Bullet')
            add_inline_text(paragraph, line[2:])
        else:
            paragraph = document.add_paragraph(style='Compact Metadata' if line.startswith('Documento') or line.startswith('Use as') else 'Normal')
            add_inline_text(paragraph, line)
        index += 1

    document.core_properties.title = 'Plano de Estudos do Concurso Público - 245 Itens'
    document.core_properties.subject = 'Catálogo consolidado de matérias e práticas de estudo'
    document.core_properties.author = 'Projeto 67 Dias'
    document.save(DOCX_PATH)
    print(f'Documento Word gerado: {DOCX_PATH}')


if __name__ == '__main__':
    if not MARKDOWN_PATH.exists():
        print(f'Arquivo de origem não encontrado: {MARKDOWN_PATH}', file=sys.stderr)
        raise SystemExit(1)
    build_document()
